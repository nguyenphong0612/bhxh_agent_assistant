import sys
import os
import re
import pdfplumber
import pypandoc
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))
from Vector_store.database import VectorDB
from Document_processing.splitter import TextSplitter
from docx import Document


class DocumentLoader:

    def load_text_file(self, file_path):

        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    # --------------------------------------------------
    # NEW: load DOCX via pypandoc (preserves heading structure as # markers)
    # --------------------------------------------------

    def load_docx_pypandoc(self, file_path):
        """
        Convert .docx → markdown using pandoc, then fix heading levels
        using python-docx ilvl (indent level from numbering XML).
        
        Pandoc flattens all headings to # level 1.
        python-docx reads ilvl: 0 = top-level (I., II.), 1 = sub (1., 2.), etc.
        We map: ilvl=0 → #, ilvl=1 → ##, ilvl=None (no numbering) → # (keep as-is).
        """
        abs_path = os.path.normpath(os.path.abspath(file_path))
        if not os.path.isfile(abs_path):
            raise FileNotFoundError(f"File not found: {abs_path}")

        # 1) python-docx: build heading text → ilvl map
        W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        doc = Document(abs_path)
        heading_map = {}  # text -> ilvl (0, 1, 2...)
        for para in doc.paragraphs:
            style = para.style.name if para.style else ""
            if 'heading' not in style.lower():
                continue
            text = para.text.strip()
            if not text:
                continue
            # Read ilvl from numPr
            pPr = para._element.pPr
            if pPr is not None:
                numPr = pPr.find(f'{W}numPr')
                if numPr is not None:
                    ilvl_el = numPr.find(f'{W}ilvl')
                    if ilvl_el is not None:
                        heading_map[text] = int(ilvl_el.get(f'{W}val'))

        # 2) pypandoc: convert to markdown
        md_text = pypandoc.convert_file(abs_path, 'markdown')

        # 3) Fix heading levels based on ilvl
        #    ilvl=0 → # (level 1), ilvl=1 → ## (level 2), etc.
        heading_re = re.compile(r'^(#{1,6})\s+(.*)')
        lines = md_text.splitlines()
        result = []
        for line in lines:
            hm = heading_re.match(line)
            if hm:
                htext = hm.group(2).strip()
                if htext:
                    for doc_text, ilvl in heading_map.items():
                        if doc_text.startswith(htext) or htext.startswith(doc_text):
                            level = ilvl + 1  # ilvl=0 → #, ilvl=1 → ##
                            line = '#' * level + ' ' + htext
                            break
            result.append(line)

        return '\n'.join(result)

    # --------------------------------------------------
    # OLD: load DOCX via python-docx (kept for reference, loses auto-numbering)
    # --------------------------------------------------

    def _is_heading_style(self, style_name):
        """Check if paragraph style indicates a heading."""
        if not style_name:
            return False
        lower_name = style_name.lower()
        return any(x in lower_name for x in ['heading', 'title', 'subtitle', 'mục', 'chương'])

    def _extract_numbering_prefix(self, para):
        """
        Attempt to extract numbering prefix from paragraph properties.
        Returns the numbering prefix if found, otherwise empty string.
        """
        try:
            # Check if paragraph has numbering properties
            pPr = para._element.pPr
            if pPr is None:
                return ""
            
            numPr = pPr.numPr
            if numPr is None:
                return ""
            
            # If numbering exists, we can't easily get the number text from python-docx
            # For now, return marker to indicate this is a numbered paragraph
            return ""  # Would require accessing document.numbering_part
        except Exception:
            return ""

    def load_docx_file(self, file_path):

        doc = Document(file_path)
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Preserve heading style information if it's a heading
                style_name = para.style.name if para.style else None
                is_heading = self._is_heading_style(style_name)
                
                # Add metadata comment only for debugging (will be stripped by metadata_extractor)
                # This ensures we don't corrupt the document structure
                if is_heading and style_name and 'heading' in style_name.lower():
                    # Just keep the text as-is, the heading structure is preserved
                    paragraphs.append(text)
                else:
                    paragraphs.append(text)

        return "\n".join(paragraphs)
    # --------------------------------------------------
    # NEW: load PDF
    # --------------------------------------------------

    def load_pdf_file(self, file_path):

        full_text = []

        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()

                if text:
                    full_text.append(text)

        return "\n".join(full_text)

    # --------------------------------------------------
    # AUTO detect file type (BEST PRACTICE)
    # --------------------------------------------------

    def load(self, file_path):

        if file_path.endswith(".txt"):
            return self.load_text_file(file_path)

        elif file_path.endswith(".docx"):
            return self.load_docx_pypandoc(file_path)
        
        elif file_path.endswith(".pdf"):
            return self.load_pdf_file(file_path)

        else:
            raise ValueError("Unsupported file format")

